from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.conf import settings
from .models import Documento, DocumentoUsuario, HistorialLecturaExamen, RecepcionDocumento
from .forms import DocumentoForm, DocumentoUsuarioForm
from users.models import User
from personal.models import Trabajador
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
from django.views.decorators.http import require_POST
from datetime import datetime


# ========== DOCUMENTOS GENERALES ==========

@login_required
def lista_documentos(request):
    if not request.user.es_admin:
        messages.error(request, 'No tienes permiso.')
        return redirect('dashboard')
    documentos = Documento.objects.all().order_by('-fecha_subida')
    return render(request, 'documentos/lista.html', {'documentos': documentos})


@login_required
def subir_documento(request):
    if not request.user.es_admin:
        return redirect('dashboard')
    form = DocumentoForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        doc = form.save(commit=False)
        doc.creado_por = request.user
        doc.save()
        messages.success(request, f'Documento "{doc.titulo}" subido correctamente.')
        return redirect('documentos:lista')
    return render(request, 'documentos/form.html', {'form': form, 'titulo': 'Subir documento PDF'})


@login_required
def editar_documento(request, pk):
    if not request.user.es_admin:
        return redirect('dashboard')
    doc = get_object_or_404(Documento, pk=pk)
    form = DocumentoForm(request.POST or None, request.FILES or None, instance=doc)
    if request.method == 'POST' and form.is_valid():
        form.save()
        messages.success(request, 'Documento actualizado.')
        return redirect('documentos:lista')
    return render(request, 'documentos/form.html', {'form': form, 'titulo': 'Editar documento'})


@login_required
def eliminar_documento(request, pk):
    if not request.user.es_admin:
        return redirect('dashboard')
    doc = get_object_or_404(Documento, pk=pk)
    doc.delete()
    messages.success(request, 'Documento eliminado.')
    return redirect('documentos:lista')


@login_required
def ver_documento(request, pk):
    doc = get_object_or_404(Documento, pk=pk, activo=True)
    return render(request, 'documentos/ver.html', {'doc': doc})


# ========== DOCUMENTOS DE USUARIO ==========

@login_required
def subir_documento_usuario(request):
    if request.method == 'POST':
        form = DocumentoUsuarioForm(request.POST, request.FILES)
        if form.is_valid():
            doc = form.save(commit=False)
            doc.usuario = request.user
            doc.save()
            messages.success(request, 'Documento subido correctamente.')
            return redirect('documentos:mis_documentos')
    else:
        form = DocumentoUsuarioForm()
    return render(request, 'documentos/subir_usuario.html', {'form': form})


@login_required
def mis_documentos(request):
    documentos = DocumentoUsuario.objects.filter(usuario=request.user)
    return render(request, 'documentos/mis_documentos.html', {'documentos': documentos})


@login_required
def revisar_documentos_usuario(request):
    if not request.user.es_admin:
        return redirect('dashboard')
    documentos = DocumentoUsuario.objects.all().order_by('-fecha_subida')
    return render(request, 'documentos/revisar_usuario.html', {'documentos': documentos})


@login_required
def cambiar_estado_documento(request, pk, estado):
    if not request.user.es_admin:
        return redirect('dashboard')
    doc = get_object_or_404(DocumentoUsuario, pk=pk)
    doc.estado = estado
    observacion = request.GET.get('observacion', '')
    if observacion:
        doc.observacion = observacion
    from django.utils import timezone
    doc.fecha_revision = timezone.now()
    doc.save()
    messages.success(request, f'Documento "{doc.titulo}" ha sido {estado}.')
    return redirect('documentos:revisar_documentos_usuario')


# ========== FUNCIONES AUXILIARES PARA DOCX ==========

def set_cell_border(cell, border_size=1):
    """Agrega bordes a una celda de tabla en DOCX"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), str(border_size))
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcPr.append(edge_el)


# ========== FORMATO F-03 EN DOCX (PLANTILLA WORD) ==========

@login_required
def generar_f03_docx(request, trabajador_pk):
    """Genera el Formato F-03 exactamente como la plantilla Word"""
    if not request.user.es_admin:
        messages.error(request, 'No tienes permiso.')
        return redirect('dashboard')
    
    # Ruta de la plantilla
    plantilla_path = os.path.join(settings.BASE_DIR, 'templates', 'formatos', 'F-03_plantilla.docx')
    
    # Verificar si la plantilla existe
    if not os.path.exists(plantilla_path):
        messages.error(request, 'No se encontró la plantilla F-03. Contacta al administrador.')
        return redirect('personal:detalle', pk=trabajador_pk)
    
    # Cargar la plantilla
    doc = Document(plantilla_path)
    
    trabajador = get_object_or_404(Trabajador, pk=trabajador_pk)
    recepciones = RecepcionDocumento.objects.filter(
        trabajador=trabajador, 
        firmado=True
    ).select_related('documento').order_by('fecha_recepcion')
    
    # Buscar la tabla en el documento (primera tabla)
    if len(doc.tables) == 0:
        messages.error(request, 'La plantilla no contiene una tabla válida.')
        return redirect('personal:detalle', pk=trabajador_pk)
    
    tabla = doc.tables[0]
    
    # Contar filas actuales
    filas_actuales = len(tabla.rows)
    
    # Agregar filas según la cantidad de documentos recibidos
    filas_necesarias = len(recepciones)
    
    # Si necesitamos más filas, las agregamos
    for i in range(filas_necesarias - (filas_actuales - 1)):
        tabla.add_row()
    
    # Llenar los datos
    for idx, rec in enumerate(recepciones):
        row_idx = idx + 1  # +1 porque la fila 0 es el encabezado
        
        if row_idx < len(tabla.rows):
            row = tabla.rows[row_idx]
            
            # Nº
            row.cells[0].text = str(idx + 1)
            
            # FECHA DE ENTREGA / RECEPCIÓN
            row.cells[1].text = rec.fecha_recepcion.strftime('%d/%m/%Y')
            
            # CÓDIGO/ NOMBRE DEL DOCUMENTO
            nombre_doc = rec.documento.titulo
            if hasattr(rec.documento, 'codigo') and rec.documento.codigo:
                nombre_doc = f"{rec.documento.codigo} - {rec.documento.titulo}"
            row.cells[2].text = nombre_doc
            
            # VERSIÓN
            version = getattr(rec.documento, 'version', 'v01')
            row.cells[3].text = version
            
            # RECIBIDO POR (Firma)
            if rec.firma_imagen:
                try:
                    firma_path = os.path.join(settings.MEDIA_ROOT, rec.firma_imagen)
                    if os.path.exists(firma_path):
                        # Limpiar párrafo existente
                        row.cells[4].paragraphs[0].clear()
                        # Agregar imagen de firma
                        run = row.cells[4].paragraphs[0].add_run()
                        run.add_picture(firma_path, width=Cm(2.5), height=Cm(1))
                    else:
                        row.cells[4].text = "Firmado digitalmente"
                except Exception:
                    row.cells[4].text = "Firmado digitalmente"
            else:
                row.cells[4].text = "Firmado digitalmente"
            
            # N° COPIAS (siempre 1)
            row.cells[5].text = "1"
            
            # DEVOLUCIÓN DE DOCUMENTO OBSOLETO (vacío)
            row.cells[6].text = ""
            
            # Aplicar bordes a las celdas
            for col in range(7):
                set_cell_border(row.cells[col])
    
    # Preparar respuesta
    filename = f"F-03_{trabajador.usuario.last_name}_{trabajador.dni}.docx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    
    return response

@login_required
def historial_documento(request, pk):
    if not request.user.es_admin:
        messages.error(request, 'No tienes permiso.')
        return redirect('dashboard')

    documento = get_object_or_404(Documento, pk=pk)

    historial = HistorialLecturaExamen.objects.filter(
        documento=documento
    ).select_related('usuario').order_by('-fecha_lectura')

    return render(request, 'documentos/historial.html', {
        'documento': documento,
        'historial': historial,
    })

@login_required
def api_usuarios(request):
    if not request.user.es_admin:
        return JsonResponse({'success': False})

    trabajadores = Trabajador.objects.select_related('usuario').all()

    usuarios = []
    for t in trabajadores:
        usuarios.append({
            'id': t.usuario.id,
            'nombre': t.usuario.username,
            'nombre_completo': t.usuario.get_full_name() or t.usuario.username
        })

    return JsonResponse({
        'success': True,
        'usuarios': usuarios
    })


@login_required
def api_historial_documento(request, pk):
    if not request.user.es_admin:
        return JsonResponse({'success': False})

    historial = HistorialLecturaExamen.objects.filter(documento_id=pk)

    data = []
    for h in historial:
        data.append({
            'usuario': h.usuario.username,
            'fechaLectura': h.fecha_lectura.strftime('%Y-%m-%d') if h.fecha_lectura else ''
        })

    return JsonResponse({
        'success': True,
        'data': data
    })


@login_required
@require_POST
def guardar_historial_documento(request, pk):
    if not request.user.es_admin:
        return JsonResponse({'success': False, 'error': 'Sin permiso'})

    try:
        body = json.loads(request.body)
        historial = body.get('historial', [])

        for item in historial:
            username = item['usuario']
            fecha = item['fechaLectura']

            user = User.objects.get(username=username)

            obj, created = HistorialLecturaExamen.objects.get_or_create(
                usuario=user,
                documento_id=pk
            )

            obj.fecha_lectura = fecha
            obj.save()

        return JsonResponse({'success': True})

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})